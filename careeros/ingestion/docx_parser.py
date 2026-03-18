"""DOCX text extraction using python-docx."""
from careeros.core.logging import get_logger

log = get_logger(__name__)


def parse_docx(file_path: str) -> str:
    """Extract raw text from DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX parsing: pip install python-docx")

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract table text
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)

    all_text = "\n".join(paragraphs + table_texts)
    log.info("docx_parser.complete", file_path=file_path, paragraphs=len(paragraphs), chars=len(all_text))
    return all_text
